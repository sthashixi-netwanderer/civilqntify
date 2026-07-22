#!/usr/bin/env python3
"""Generate the Material Quantification Output Format table for Section 3.4.2 as .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Table_3.4.2_Material_Quantification_Output.docx")

# ── Table data ────────────────────────────────────────────────────────────────
# (category, #, param, formula/method, source, unit, output/example)
ROWS = [
    # ── Section A: Structural Element Input ────────────────────────────────────
    ("A. Structural Element Input", None, None, None, None, None, None),
    (None, 1, "Element type",
     "User selects from predefined list",
     "User input",
     "—",
     "Slab / Beam / Column / Foundation / Wall / Staircase"),
    (None, 2, "Length (L)",
     "Measured per IS 1200 rules",
     "User input",
     "m",
     "e.g., 6.00 m"),
    (None, 3, "Width (W) / Breadth (B)",
     "Measured per IS 1200 rules",
     "User input",
     "m",
     "e.g., 3.00 m"),
    (None, 4, "Thickness (T) / Height (H)",
     "Measured per IS 1200 rules",
     "User input",
     "m",
     "e.g., 0.15 m (slab), 0.45 m (beam depth)"),
    (None, 5, "Number of identical elements",
     "Count of repeating elements",
     "User input",
     "—",
     "e.g., 12"),
    (None, 6, "Deductions — openings > 0.1 m²",
     "V_ded = L_open × W_open × T",
     "IS 1200, §3.4.2 description",
     "m³",
     "e.g., 0.30 m² window → 0.30 × 0.15 = 0.045 m³"),

    # ── Section B: Volume Calculations ────────────────────────────────────────
    ("B. Volume Calculations", None, None, None, None, None, None),
    (None, 7, "Gross volume per element",
     "V_gross = L × W × T  (slab)\nV_gross = L × W × H  (beam — face-to-face of columns)\nV_gross = L × W × H  (column — full height)",
     "§3.4.2 description",
     "m³",
     "e.g., 6.00 × 3.00 × 0.15 = 2.700 m³"),
    (None, 8, "Total gross volume (all elements)",
     "V_total_gross = Σ (V_gross × N)",
     "Summation across all element types",
     "m³",
     "e.g., (2.700 × 12) + (1.350 × 8) = 43.20 m³"),
    (None, 9, "Total deductions",
     "V_deductions = Σ (V_ded per opening)",
     "IS 1200 measurement rules",
     "m³",
     "e.g., 0.045 × 6 = 0.270 m³"),
    (None, 10, "Net wet volume",
     "V_wet = V_total_gross − V_deductions",
     "Calculated",
     "m³",
     "e.g., 43.20 − 0.27 = 42.93 m³"),
    (None, 11, "Dry volume conversion factor",
     "1.54 (accounts for ~35% air void elimination)",
     "§3.4.2 description",
     "—",
     "1.54 (fixed)"),
    (None, 12, "Total dry volume",
     "V_dry = V_wet × 1.54",
     "§3.4.2 description",
     "m³",
     "e.g., 42.93 × 1.54 = 66.11 m³"),

    # ── Section C: Mix Design Proportions (per m³) ────────────────────────────
    ("C. Mix Design Proportions (per m³ of concrete)", None, None, None, None, None, None),
    (None, 13, "Cement content",
     "From Mix Design Module (§3.4.1)",
     "ACI Table 5.3.4 / IS Fig. 1",
     "kg/m³",
     "e.g., 350 kg/m³"),
    (None, 14, "Water content",
     "From Mix Design Module (§3.4.1)",
     "ACI Table 5.3.3 / IS Table 4",
     "kg/m³",
     "e.g., 180 kg/m³"),
    (None, 15, "Fine aggregate content",
     "From Mix Design Module (§3.4.1)",
     "ACI Table 5.3.6 / IS Table 5",
     "kg/m³",
     "e.g., 680 kg/m³"),
    (None, 16, "Coarse aggregate content",
     "From Mix Design Module (§3.4.1)",
     "ACI Table 5.3.6 / IS Table 5",
     "kg/m³",
     "e.g., 1,100 kg/m³"),
    (None, 17, "SCM content (if applicable)",
     "From Mix Design Module (§3.4.1)",
     "§4.1(p), IS Table 9",
     "kg/m³",
     "e.g., 70 kg/m³ (fly ash)"),
    (None, 18, "Admixture content (if applicable)",
     "From Mix Design Module (§3.4.1)",
     "§6.3, IS Annex G",
     "kg/m³ or L/m³",
     "e.g., 1.75 L/m³"),

    # ── Section D: Total Material Quantities ──────────────────────────────────
    ("D. Total Material Quantities", None, None, None, None, None, None),
    (None, 19, "Total cement",
     "M_cement = Cement content × V_dry",
     "§3.4.2 description",
     "kg",
     "e.g., 350 × 66.11 = 23,139 kg"),
    (None, 20, "Total water",
     "M_water = Water content × V_dry",
     "§3.4.2 description",
     "kg (litres)",
     "e.g., 180 × 66.11 = 11,900 kg"),
    (None, 21, "Total fine aggregate",
     "M_fine = Fine agg. content × V_dry",
     "§3.4.2 description",
     "kg",
     "e.g., 680 × 66.11 = 44,955 kg"),
    (None, 22, "Total coarse aggregate",
     "M_coarse = Coarse agg. content × V_dry",
     "§3.4.2 description",
     "kg",
     "e.g., 1,100 × 66.11 = 72,721 kg"),
    (None, 23, "Total SCM",
     "M_SCM = SCM content × V_dry",
     "§3.4.2 description",
     "kg",
     "e.g., 70 × 66.11 = 4,628 kg"),
    (None, 24, "Total admixture",
     "M_admix = Admixture content × V_dry",
     "§3.4.2 description",
     "kg or L",
     "e.g., 1.75 × 66.11 = 115.7 L"),

    # ── Section E: Procurement Conversion ──────────────────────────────────────
    ("E. Procurement Conversion (Ghana Standard Units)", None, None, None, None, None, None),
    (None, 25, "Cement — 50 kg bags",
     "Bags = M_cement / 50",
     "§3.4.2 description",
     "bags",
     "e.g., 23,139 / 50 = 463 bags"),
    (None, 26, "Fine aggregate — truckloads",
     "Trucks = M_fine / (truck capacity × loose density)",
     "Local truck capacity (typically 10–14 m³ loose)",
     "truckloads",
     "e.g., ~45,000 kg ÷ 1,500 kg/m³ = 30 m³ ≈ 3 trucks"),
    (None, 27, "Coarse aggregate — truckloads",
     "Trucks = M_coarse / (truck capacity × loose density)",
     "Local truck capacity (typically 10–14 m³ loose)",
     "truckloads",
     "e.g., ~72,700 kg ÷ 1,450 kg/m³ = 50 m³ ≈ 4 trucks"),
    (None, 28, "Water — tankers / drums",
     "Tankers = M_water / tanker capacity",
     "Local tanker capacity (typically 5,000–10,000 L)",
     "tankers",
     "e.g., 11,900 L ≈ 1–2 tankers"),

    # ── Section F: Wastage & Final Order ──────────────────────────────────────
    ("F. Wastage Allowance & Final Order Quantity", None, None, None, None, None, None),
    (None, 29, "Wastage factor",
     "User-specified or default (2–5%)",
     "§3.4.2 description",
     "%",
     "e.g., 3%"),
    (None, 30, "Wastage allowance — cement",
     "W_cement = M_cement × (wastage / 100)",
     "§3.4.2 description",
     "kg",
     "e.g., 23,139 × 0.03 = 694 kg"),
    (None, 31, "Wastage allowance — fine aggregate",
     "W_fine = M_fine × (wastage / 100)",
     "§3.4.2 description",
     "kg",
     "e.g., 44,955 × 0.03 = 1,349 kg"),
    (None, 32, "Wastage allowance — coarse aggregate",
     "W_coarse = M_coarse × (wastage / 100)",
     "§3.4.2 description",
     "kg",
     "e.g., 72,721 × 0.03 = 2,182 kg"),
    (None, 33, "Wastage allowance — water",
     "W_water = M_water × (wastage / 100)",
     "§3.4.2 description",
     "kg (litres)",
     "e.g., 11,900 × 0.03 = 357 L"),
    (None, 34, "Final order quantity — cement",
     "M_cement + W_cement → round up to bags",
     "§3.4.2 description",
     "bags (50 kg)",
     "e.g., 23,833 kg ÷ 50 = 477 bags"),
    (None, 35, "Final order quantity — fine aggregate",
     "M_fine + W_fine → convert to truckloads",
     "§3.4.2 description",
     "truckloads",
     "e.g., 46,304 kg → ~31 m³ ≈ 3 trucks"),
    (None, 36, "Final order quantity — coarse aggregate",
     "M_coarse + W_coarse → convert to truckloads",
     "§3.4.2 description",
     "truckloads",
     "e.g., 74,903 kg → ~52 m³ ≈ 4 trucks"),
    (None, 37, "Final order quantity — water",
     "M_water + W_water",
     "§3.4.2 description",
     "litres",
     "e.g., 12,257 L"),
    (None, 38, "Final order quantity — SCM",
     "M_SCM + W_SCM",
     "§3.4.2 description",
     "kg",
     "e.g., 4,628 + 139 = 4,767 kg"),
    (None, 39, "Final order quantity — admixture",
     "M_admix + W_admix",
     "§3.4.2 description",
     "L",
     "e.g., 115.7 + 3.5 = 119.2 L"),
]

HEADERS = ["#", "Parameter", "Formula / Method", "Source / Reference",
           "Unit", "Output / Example"]

# ── Colour palette ────────────────────────────────────────────────────────────
HEADER_BG = "1B4332"   # dark green (differentiated from Mix Design module)
HEADER_FG = "FFFFFF"
CAT_BG    = "D8F3DC"   # light green for category rows
CAT_FG    = "1B4332"
ALT_BG    = "F0FAF4"   # subtle green tint
WHITE_BG  = "FFFFFF"
BORDER    = "74C69D"   # green-grey borders

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, color=BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val_map = {"top": "top", "center": "center", "bottom": "bottom"}
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val_map.get(align, "center")}"/>')
    tcPr.append(vAlign)

def add_cell_text(cell, text, bold=False, font_size=8, font_name="Calibri",
                  color="000000", alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def build_document():
    doc = Document()

    # ── Page setup — A4 landscape ─────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = 1
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Table 3.4.2 — Material Quantification Output Format", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    # ── Subtitle ──────────────────────────────────────────────────────────────
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run(
        "Module: Material Quantification (§3.4.2)  |  "
        "Converts per-m³ mix proportions into total project quantities  |  "
        "Wet-to-dry conversion factor: 1.54"
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")

    # ── Table ─────────────────────────────────────────────────────────────────
    ncols = 6
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Column widths — #, Param, Formula, Source, Unit, Example
    col_widths = [Cm(0.8), Cm(4.0), Cm(7.0), Cm(4.5), Cm(2.2), Cm(5.7)]

    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, h in enumerate(HEADERS):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell, HEADER_BG)
        set_cell_vertical_alignment(cell, "center")
        add_cell_text(cell, h, bold=True, font_size=8.5, font_name="Calibri",
                      color=HEADER_FG, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    tr = hdr_row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="380" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    # ── Data rows ─────────────────────────────────────────────────────────────
    data_row_idx = 0
    for row_data in ROWS:
        cat = row_data[0]
        if cat is not None:
            row = table.add_row()
            start_cell = row.cells[0]
            end_cell = row.cells[ncols - 1]
            start_cell.merge(end_cell)
            merged = row.cells[0]
            set_cell_shading(merged, CAT_BG)
            set_cell_borders(merged, BORDER)
            set_cell_vertical_alignment(merged, "center")
            add_cell_text(merged, cat, bold=True, font_size=9, font_name="Calibri",
                          color=CAT_FG, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="300" w:hRule="atLeast"/>')
            trPr.append(trHeight)
            data_row_idx = 0
        else:
            _, num, param, formula, source, unit, example = row_data
            row = table.add_row()
            bg = ALT_BG if data_row_idx % 2 == 0 else WHITE_BG
            cells = row.cells
            texts = [str(num), param, formula, source, unit, example]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            bolds = [True, True, False, False, False, False]
            for i in range(ncols):
                c = cells[i]
                set_cell_shading(c, bg)
                set_cell_borders(c, BORDER)
                set_cell_vertical_alignment(c, "center")
                add_cell_text(c, texts[i], bold=bolds[i], font_size=7.5,
                              font_name="Calibri", color="1A1A1A", alignment=aligns[i])
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>')
            trPr.append(trHeight)
            data_row_idx += 1

    # ── Notes ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    run = notes_heading.add_run("Notes:")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    notes = [
        "The wet-to-dry conversion factor of 1.54 accounts for the ~35% air void space in dry aggregates that is eliminated upon mixing and compaction.",
        "Deductions for openings exceeding 0.1 m² are made in accordance with IS 1200 measurement rules.",
        "Beam volumes are measured face-to-face of columns; column volumes are measured full height.",
        "Cement is converted to 50 kg bags (standard unit in Ghana); aggregates are converted to truckloads based on local truck capacity and loose bulk density.",
        "The wastage factor of 2–5% accounts for spillage, formwork bulging, and pumping losses. The user may adjust this based on site conditions.",
        "All mix proportions (Section C) are sourced from the Mix Design Module (§3.4.1). The quantification module does not independently calculate w/cm or aggregate ratios.",
    ]
    for n in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(n)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("333333")

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")

if __name__ == "__main__":
    build_document()
