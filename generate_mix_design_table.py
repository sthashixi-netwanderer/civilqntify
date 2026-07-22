#!/usr/bin/env python3
"""Generate the Input Parameters table for the Concrete Mix Design Module as .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Table_3.4.1_Input_Parameters_Mix_Design.docx")

# ── Table data ────────────────────────────────────────────────────────────────
# (category, #, param, desc, aci_ref, is_ref, unit, values)
# category=None means it's a data row; category=string means it's a section header
ROWS = [
    ("General", None, None, None, None, None, None, None),
    (None, 1, "Design method",
     "Standard to use for proportioning",
     "§5.3",
     "§5–§9",
     "—",
     "ACI 211.1-22 / IS 10262:2019"),
    (None, 2, "Target compressive strength",
     "Characteristic strength at 28 days",
     "§4.7.4, Table 4.7.4.1",
     "§4.2, Table 1",
     "MPa (N/mm²) or psi",
     "User-specified"),
    (None, 3, "Standard deviation",
     "For target mean strength calculation",
     "§4.7.4.2–4.7.4.4",
     "§4.2.1, Table 2",
     "MPa or psi",
     "Default from Table 2 or user-supplied"),

    ("Aggregates", None, None, None, None, None, None, None),
    (None, 4, "Nominal maximum aggregate size",
     "Largest aggregate permitted",
     "§4.3.2, §5.3.2",
     "§4.1(c), Table 3, Table 4",
     "mm or in.",
     "10, 20, 40 mm / 3/8, 1/2, 3/4, 1, 1-1/2, 2, 3 in."),
    (None, 5, "Fine aggregate grading zone",
     "Sand grading classification",
     "—",
     "IS 383, Table 5",
     "—",
     "Zone I, Zone II, Zone III, Zone IV"),
    (None, 6, "Coarse aggregate shape",
     "Particle shape of coarse aggregate",
     "§4.3",
     "§4.1(j)",
     "—",
     "Angular, Sub-angular, Rounded, Manufactured"),
    (None, 7, "Fine aggregate type",
     "Source of fine aggregate",
     "§4.3",
     "§4.1(k)",
     "—",
     "Natural sand, Crushed stone sand, Manufactured sand, Mixed sand"),
    (None, 8, "Fineness modulus of fine aggregate",
     "From sieve analysis (ASTM C136)",
     "§4.3.5, Table 5.3.6",
     "—",
     "—",
     "2.40–3.00"),
    (None, 9, "Specific gravity — coarse aggregate (SSD)",
     "Bulk specific gravity at SSD condition",
     "§4.3.7, ASTM C127",
     "§4.1, IS 2386(Part 3)",
     "—",
     "Typically 2.6–2.8"),
    (None, 10, "Specific gravity — fine aggregate (SSD)",
     "Bulk specific gravity at SSD condition",
     "§4.3.7, ASTM C128",
     "§4.1, IS 2386(Part 3)",
     "—",
     "Typically 2.5–2.7"),
    (None, 11, "Specific gravity — cement",
     "Specific gravity of cement used",
     "§4.7.5",
     "§4.1",
     "—",
     "3.15 (OPC), or actual value"),
    (None, 12, "Dry-rodded unit weight (DRUW)",
     "Oven-dry aggregate compacted by rodding (ASTM C29)",
     "§4.3.6, Table 5.3.6",
     "—",
     "kg/m³ or lb/ft³",
     "From lab test"),
    (None, 13, "Absorption — coarse aggregate",
     "Moisture absorption capacity",
     "§4.7.8.4",
     "IS 2386(Part 3)",
     "%",
     "From lab test"),
    (None, 14, "Absorption — fine aggregate",
     "Moisture absorption capacity",
     "§4.7.8.4",
     "IS 2386(Part 3)",
     "%",
     "From lab test"),

    ("Workability", None, None, None, None, None, None, None),
    (None, 15, "Target slump",
     "Concrete consistency requirement",
     "§5.3.1, Table 5.3.1",
     "§5.3, Table 4",
     "mm or in.",
     "25–150 mm / 1–7 in."),
    (None, 16, "Slump adjustment for admixtures",
     "Water content correction when WRA/HRWRA used",
     "Table 5.3.3.1",
     "§5.3",
     "—",
     "WRA: −5%, HRWRA: −12%"),

    ("Water & Air", None, None, None, None, None, None, None),
    (None, 17, "Air entrainment",
     "Whether air-entraining admixture is used",
     "§4.6, Table 5.3.3",
     "§5.2, Table 3",
     "Yes/No",
     "Yes or No"),
    (None, 18, "Target air content",
     "Total air in fresh concrete",
     "Table 5.3.3, §4.6.2",
     "Table 3 (entrapped only)",
     "%",
     "Non-air-entrained: 0.3–3.0%; Air-entrained: 3.5–7.5%"),
    (None, 19, "Water content (initial estimate)",
     "From slump/aggregate size lookup",
     "Table 5.3.3",
     "Table 4",
     "kg/m³ or lb/yd³",
     "Non-air-entrained vs. air-entrained tables"),
    (None, 20, "Water content adjustments",
     "Corrections for temperature, shape, admixtures",
     "Table 5.3.3.1",
     "§5.3",
     "%",
     "Per-condition adjustments"),

    ("Cementitious Materials", None, None, None, None, None, None, None),
    (None, 21, "w/cm ratio (strength-based)",
     "From strength relationship or lookup",
     "Table 5.3.4",
     "Fig. 1",
     "—",
     "Typically 0.30–0.80"),
    (None, 22, "w/cm ratio (durability-based)",
     "Maximum from exposure class",
     "Table 4.7.3a–4.7.3d",
     "IS 456 Table 3 & 5",
     "—",
     "Varies by exposure class"),
    (None, 23, "Minimum cementitious content",
     "From durability requirements",
     "§4.7.3, Tables 4.7.3a–d",
     "IS 456 Table 3 & 5",
     "kg/m³ or lb/yd³",
     "Varies by exposure class"),
    (None, 24, "Maximum cementitious content",
     "Specification or IS 456 limit",
     "Per project spec",
     "IS 456 §6",
     "kg/m³ or lb/yd³",
     "Typically ≤450 kg/m³"),
    (None, 25, "Cement type",
     "ASTM/IS cement classification",
     "§4.7.5",
     "§4.1(b)",
     "—",
     "ASTM C150 Type I–V / IS 269, IS 8112, IS 12269"),

    ("Supplementary Cementitious Materials (SCMs)", None, None, None, None, None, None, None),
    (None, 26, "SCM type",
     "Fly ash, slag, silica fume, etc.",
     "§4.7.6, Chapter 7",
     "§4.1(p), Table 9",
     "—",
     "Fly ash (Class F/C), GGBS, Silica fume, Metakaolin"),
    (None, 27, "SCM replacement percentage",
     "% of total cementitious by mass",
     "§4.7.6, Table 4.7.3.2",
     "Table 9",
     "%",
     "Fly ash: 15–30%; GGBS: 25–50%; Silica fume: 5–10%"),
    (None, 28, "Specific gravity of SCM",
     "For absolute volume calculation",
     "§4.7.6",
     "—",
     "—",
     "Fly ash: ~2.3; GGBS: ~2.9; Silica fume: ~2.2"),

    ("Chemical Admixtures", None, None, None, None, None, None, None),
    (None, 29, "Admixture type",
     "WRA, HRWRA, air-entraining, retarder, accelerator",
     "§6.3, §4.5",
     "§4.1(n), Annex G",
     "—",
     "ASTM C494 Types A–G / IS 9103"),
    (None, 30, "Admixture dosage",
     "Per manufacturer recommendation",
     "§6.3",
     "Annex G",
     "mL/100 kg cementitious",
     "From lab test or manufacturer data"),
    (None, 31, "Admixture specific gravity",
     "For absolute volume calculation",
     "§4.5",
     "—",
     "—",
     "Typically 1.0–1.2"),

    ("Environmental / Exposure", None, None, None, None, None, None, None),
    (None, 32, "Exposure class (ACI)",
     "S, F, W, C categories per ACI 318 Ch. 19",
     "Tables 4.7.3a–4.7.3d",
     "—",
     "—",
     "S0–S3, F0–F3, W0–W2, C0–C2"),
    (None, 33, "Exposure condition (IS)",
     "Per IS 456 Table 3 and Table 5",
     "—",
     "§4.1(e), IS 456 Table 3 & 5",
     "—",
     "Moderate, Severe, Very Severe, etc."),

    ("Field Adjustments", None, None, None, None, None, None, None),
    (None, 34, "Moisture content — coarse aggregate",
     "Total moisture on aggregate stockpile",
     "§4.7.8.3",
     "—",
     "%",
     "From site test"),
    (None, 35, "Moisture content — fine aggregate",
     "Total moisture on aggregate stockpile",
     "§4.7.8.3",
     "—",
     "%",
     "From site test"),
    (None, 36, "Free moisture content — coarse",
     "Free water available beyond SSD",
     "§4.7.8.5",
     "—",
     "%",
     "Calculated: %total − absorption"),
    (None, 37, "Free moisture content — fine",
     "Free water available beyond SSD",
     "§4.7.8.5",
     "—",
     "%",
     "Calculated: %total − absorption"),
]

HEADERS = ["#", "Parameter", "Description", "ACI 211.1-22\nReference",
           "IS 10262:2019\nReference", "Unit", "Values / Options"]

# ── Colour palette ────────────────────────────────────────────────────────────
HEADER_BG = "1B3A5C"   # dark navy
HEADER_FG = "FFFFFF"   # white text
CAT_BG    = "D6E4F0"   # light blue for category rows
CAT_FG    = "1B3A5C"   # dark navy text
ALT_BG    = "F2F6FA"   # subtle alternating row tint
WHITE_BG  = "FFFFFF"
BORDER    = "8EA9C1"   # medium blue-grey borders

def set_cell_shading(cell, hex_color):
    """Apply background shading to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, color=BORDER):
    """Set thin borders on all four sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val_map = {"top": "top", "center": "center", "bottom": "bottom"}
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val_map.get(align, "center")}"/>')
    tcPr.append(vAlign)

def add_cell_text(cell, text, bold=False, font_size=8, font_name="Calibri",
                  color="000000", alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Write text into a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color)
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def build_document():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(29.7)    # A4 landscape
    section.page_height = Cm(21.0)
    section.orientation = 1          # landscape
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Table 3.4.1 — Input Parameters for Concrete Mix Design Module", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    # ── Subtitle / source note ────────────────────────────────────────────────
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(
        "Source: ACI PRC-211.1-22 (Selecting Proportions for Normal-Density and High-Density Concrete) "
        "and IS 10262:2019 (Concrete Mix Proportioning — Guidelines, Second Revision)."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")

    # ── Table ─────────────────────────────────────────────────────────────────
    # Columns: #, Parameter, Description, ACI Ref, IS Ref, Unit, Values
    ncols = 7
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Column widths (total ~26.7 cm usable)
    col_widths = [Cm(0.8), Cm(4.2), Cm(4.5), Cm(3.5), Cm(3.5), Cm(2.5), Cm(5.7)]

    # Set column widths
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, h in enumerate(HEADERS):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell, HEADER_BG)
        set_cell_vertical_alignment(cell, "center")
        add_cell_text(cell, h, bold=True, font_size=8, font_name="Calibri",
                      color=HEADER_FG, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Set header row height
    tr = hdr_row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    # ── Data rows ─────────────────────────────────────────────────────────────
    data_row_idx = 0  # for alternating shading
    for row_data in ROWS:
        cat = row_data[0]
        if cat is not None:
            # Category header row — merged across all columns
            row = table.add_row()
            # Merge all cells
            start_cell = row.cells[0]
            end_cell = row.cells[ncols - 1]
            start_cell.merge(end_cell)
            merged = row.cells[0]
            set_cell_shading(merged, CAT_BG)
            set_cell_borders(merged, BORDER)
            set_cell_vertical_alignment(merged, "center")
            add_cell_text(merged, cat, bold=True, font_size=9, font_name="Calibri",
                          color=CAT_FG, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            # Set row height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="300" w:hRule="atLeast"/>')
            trPr.append(trHeight)
            data_row_idx = 0  # reset alternation per category
        else:
            _, num, param, desc, aci, is_ref, unit, values = row_data
            row = table.add_row()
            bg = ALT_BG if data_row_idx % 2 == 0 else WHITE_BG
            cells = row.cells
            texts = [str(num), param, desc, aci, is_ref, unit, values]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.LEFT]
            bolds = [True, True, False, False, False, False, False]
            for i in range(ncols):
                c = cells[i]
                set_cell_shading(c, bg)
                set_cell_borders(c, BORDER)
                set_cell_vertical_alignment(c, "center")
                add_cell_text(c, texts[i], bold=bolds[i], font_size=7.5,
                              font_name="Calibri", color="1A1A1A", alignment=aligns[i])
            # Set row height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="360" w:hRule="atLeast"/>')
            trPr.append(trHeight)
            data_row_idx += 1

    # ── Notes paragraph ───────────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    notes_heading = doc.add_paragraph()
    run = notes_heading.add_run("Notes:")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    notes = [
        "ACI 211.1-22 uses English units (lb/yd³, in.) in its primary tables; the module should convert to metric (kg/m³, mm) as needed.",
        "IS 10262:2019 uses SI units (kg/m³, mm, MPa) natively.",
        "Both methods share the absolute volume principle: the sum of absolute volumes of all components (cement, water, aggregates, air, SCMs, admixtures) must equal 1.0 m³.",
        "Parameters 34–37 are only needed for Step 9 (moisture adjustments) in the field — they are not used during the initial design proportioning.",
        "For high-strength concrete (M65 and above), additional parameters apply per IS 10262 Section 3 (e.g., Table 8 w/cm for HRWRA mixes).",
    ]
    for n in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(n)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("333333")

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")

if __name__ == "__main__":
    build_document()
