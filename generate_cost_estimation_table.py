#!/usr/bin/env python3
"""Generate Sample Cost Estimation Output for a 100 m³ Project — §3.4.3 as .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Table_3.4.3_Cost_Estimation_Output.docx")

# ── Colour palette ────────────────────────────────────────────────────────────
HEADER_BG = "7B2D26"   # deep maroon (differentiated from previous modules)
HEADER_FG = "FFFFFF"
CAT_BG    = "F2E0D9"   # warm rose for category rows
CAT_FG    = "7B2D26"
ALT_BG    = "FDF5F3"   # subtle warm tint
WHITE_BG  = "FFFFFF"
TOTAL_BG  = "7B2D26"   # total row — maroon background
TOTAL_FG  = "FFFFFF"
SUBTOTAL_BG = "E8D5CE" # subtotal rows
BORDER    = "C4A49A"   # warm brown borders

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_borders(cell, color=BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val_map = {"top": "top", "center": "center", "bottom": "bottom"}
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val_map.get(align, "center")}"/>')
    tcPr.append(vAlign)

def add_cell_text(cell, text, bold=False, font_size=8, font_name="Calibri",
                  color="000000", alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

# ── Table structure ───────────────────────────────────────────────────────────
# Columns: #, Item, Quantity, Unit, Unit Price (GHS), Subtotal (GHS)
HEADERS = ["#", "Item Description", "Quantity", "Unit", "Unit Price\n(GHS)", "Subtotal\n(GHS)"]
NCOLS = 6

# Row data: (type, values...)
# type: "cat" = category header, "row" = data row, "sub" = subtotal, "total" = grand total
# For "row": (#, item, qty, unit, unit_price, subtotal)
# For "sub": (label, subtotal_value)
# For "total": (label, total_value)

TABLE_DATA = [
    # ── Section A: Material Costs ─────────────────────────────────────────────
    ("cat", "A. Material Costs"),

    ("row", ("A.1", "Cement — Dangote 42.5R (50 kg bags)", "735", "bags", "95.00", "69,825.00")),
    ("row", ("A.2", "Fine aggregate — River sand", "6", "truckloads\n(14 m³ each)", "550.00", "3,300.00")),
    ("row", ("A.3", "Coarse aggregate — 20 mm crushed stone", "9", "truckloads\n(14 m³ each)", "650.00", "5,850.00")),
    ("row", ("A.4", "Reinforcement steel — 12 mm bars", "48", "lengths\n(12 m each)", "95.00", "4,560.00")),
    ("row", ("A.5", "Reinforcement steel — 20 mm bars", "24", "lengths\n(12 m each)", "210.00", "5,040.00")),
    ("row", ("A.6", "Water (municipal supply)", "12.5", "m³", "8.50", "106.25")),

    ("sub", ("Total Material Cost (A)", "88,681.25")),

    # ── Section B: Transport / Haulage ────────────────────────────────────────
    ("cat", "B. Transport / Haulage Costs"),

    ("row", ("B.1", "Cement haulage — Dangote depot to site\n(12 km, multiplier: 1.5×)", "735", "bags", "1.50\n(per bag)", "1,102.50")),
    ("row", ("B.2", "Sand haulage — quarry to site\n(8 km, multiplier: 1.0×)", "6", "truckloads", "350.00\n(per trip)", "2,100.00")),
    ("row", ("B.3", "Aggregate haulage — quarry to site\n(8 km, multiplier: 1.0×)", "9", "truckloads", "400.00\n(per trip)", "3,600.00")),
    ("row", ("B.4", "Steel haulage — market to site\n(5 km, multiplier: 1.0×)", "72", "lengths", "5.00\n(per piece)", "360.00")),

    ("sub", ("Total Haulage Cost (B)", "7,162.50")),

    # ── Section C: Workmanship ────────────────────────────────────────────────
    ("cat", "C. Workmanship (Concrete Works)"),

    ("row", ("C.1", "Formwork — fabrication, erection & stripping", "100", "m³ placed", "45.00", "4,500.00")),
    ("row", ("C.2", "Reinforcement — cutting, bending & fixing", "1.44", "tonnes", "1,800.00", "2,592.00")),
    ("row", ("C.3", "Concrete — mixing, placing & compaction", "100", "m³ placed", "55.00", "5,500.00")),
    ("row", ("C.4", "Curing", "100", "m³ placed", "8.00", "800.00")),

    ("sub", ("Total Workmanship Cost (C)", "13,392.00")),

    # ── Section D: Contingency ────────────────────────────────────────────────
    ("cat", "D. Contingency Allowance"),

    ("row", ("D.1", "Material price fluctuation allowance (5% of A)", "—", "—", "—", "4,434.06")),
    ("row", ("D.2", "Unforeseen quantity variations (3% of A+B+C)", "—", "—", "—", "3,341.48")),

    ("sub", ("Total Contingency (D)", "7,775.54")),

    # ── Section E: Summary ────────────────────────────────────────────────────
    ("cat", "E. Project Cost Summary"),

    ("sub", ("A — Material Costs", "88,681.25")),
    ("sub", ("B — Transport / Haulage", "7,162.50")),
    ("sub", ("C — Workmanship", "13,392.00")),
    ("sub", ("D — Contingency", "7,775.54")),
    ("total", ("ESTIMATED TOTAL COST FOR 100 m³ CONCRETE WORKS", "117,011.29")),
    ("total", ("Cost per m³ of concrete", "1,170.11")),
]


def build_document():
    doc = Document()

    # ── Page setup — A4 landscape ─────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = 1
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Table 3.4.3 — Sample Cost Estimation Output for a 100 m³ Project", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    # ── Subtitle ──────────────────────────────────────────────────────────────
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(2)
    run = note.add_run(
        "Module: Cost Estimation (§3.4.3)  |  "
        "Deterministic quantity × unit-price method  |  "
        "Currency: Ghana Cedis (GHS)  |  "
        "Concrete: M25 grade, 100 m³ structural volume"
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")

    # ── Project parameters mini-table ─────────────────────────────────────────
    params = doc.add_paragraph()
    params.paragraph_format.space_after = Pt(6)
    run = params.add_run(
        "Project Parameters:  "
        "Grade: M25 (w/cm = 0.45)  |  "
        "Cement: Dangote 42.5R  |  "
        "Mix: 350 kg/m³ cement, 680 kg/m³ sand, 1,100 kg/m³ aggregate  |  "
        "Haulage distance: 8–12 km  |  "
        "Wastage: 5%  |  "
        "Contingency: 5–8%"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("444444")

    # ── Table ─────────────────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=NCOLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Column widths: #, Item, Qty, Unit, Unit Price, Subtotal
    col_widths = [Cm(0.8), Cm(7.5), Cm(2.2), Cm(2.8), Cm(2.8), Cm(3.2)]

    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, h in enumerate(HEADERS):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell, HEADER_BG)
        set_cell_vertical_alignment(cell, "center")
        add_cell_text(cell, h, bold=True, font_size=8.5, font_name="Calibri",
                      color=HEADER_FG, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    tr = hdr_row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for row_data in TABLE_DATA:
        rtype = row_data[0]

        if rtype == "cat":
            row = table.add_row()
            start_cell = row.cells[0]
            end_cell = row.cells[NCOLS - 1]
            start_cell.merge(end_cell)
            merged = row.cells[0]
            set_cell_shading(merged, CAT_BG)
            set_cell_borders(merged, BORDER)
            set_cell_vertical_alignment(merged, "center")
            add_cell_text(merged, row_data[1], bold=True, font_size=9, font_name="Calibri",
                          color=CAT_FG, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="320" w:hRule="atLeast"/>')
            trPr.append(trHeight)

        elif rtype == "row":
            num, item, qty, unit, price, subtotal = row_data[1]
            row = table.add_row()
            # Determine shading — alternate within category
            bg = ALT_BG
            cells = row.cells
            texts = [num, item, qty, unit, price, subtotal]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
            bolds = [False, False, False, False, False, False]
            for i in range(NCOLS):
                c = cells[i]
                set_cell_shading(c, bg)
                set_cell_borders(c, BORDER)
                set_cell_vertical_alignment(c, "center")
                add_cell_text(c, texts[i], bold=bolds[i], font_size=8,
                              font_name="Calibri", color="1A1A1A", alignment=aligns[i])
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="420" w:hRule="atLeast"/>')
            trPr.append(trHeight)

        elif rtype == "sub":
            label, value = row_data[1]
            row = table.add_row()
            # Merge first 5 columns for label, last column for value
            start_cell = row.cells[0]
            end_cell = row.cells[NCOLS - 2]
            start_cell.merge(end_cell)
            merged = row.cells[0]
            last_cell = row.cells[NCOLS - 1]

            set_cell_shading(merged, SUBTOTAL_BG)
            set_cell_borders(merged, BORDER)
            set_cell_vertical_alignment(merged, "center")
            add_cell_text(merged, label, bold=True, font_size=8.5, font_name="Calibri",
                          color=HEADER_BG, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            set_cell_shading(last_cell, SUBTOTAL_BG)
            set_cell_borders(last_cell, BORDER)
            set_cell_vertical_alignment(last_cell, "center")
            add_cell_text(last_cell, value, bold=True, font_size=8.5, font_name="Calibri",
                          color=HEADER_BG, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="340" w:hRule="atLeast"/>')
            trPr.append(trHeight)

        elif rtype == "total":
            label, value = row_data[1]
            row = table.add_row()
            start_cell = row.cells[0]
            end_cell = row.cells[NCOLS - 2]
            start_cell.merge(end_cell)
            merged = row.cells[0]
            last_cell = row.cells[NCOLS - 1]

            set_cell_shading(merged, TOTAL_BG)
            set_cell_borders(merged, TOTAL_BG)
            set_cell_vertical_alignment(merged, "center")
            add_cell_text(merged, label, bold=True, font_size=10, font_name="Calibri",
                          color=TOTAL_FG, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            set_cell_shading(last_cell, TOTAL_BG)
            set_cell_borders(last_cell, TOTAL_BG)
            set_cell_vertical_alignment(last_cell, "center")
            add_cell_text(last_cell, f"GHS {value}", bold=True, font_size=10,
                          font_name="Calibri", color=TOTAL_FG,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>')
            trPr.append(trHeight)

    # ── Notes ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    run = notes_heading.add_run("Notes:")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(HEADER_BG)

    notes = [
        "All prices are in Ghana Cedis (GHS) and reflect current market rates. The JSON price database can be updated without modifying source code.",
        "Haulage multipliers: 0–5 km = 1.0×, 6–15 km = 1.5×, 16–50 km = 2.0×. Distance is measured from the nearest supplier depot.",
        "Cement is priced per 50 kg bag (Ghana standard). Aggregates are priced per truckload (≈14 m³ loose volume). Steel is priced per 12 m length.",
        "Contingency D.1 covers material price fluctuations (default 5% of material cost). Contingency D.2 covers unforeseen quantity variations (default 3% of subtotal A+B+C). Both are user-adjustable.",
        "Workmanship rates (Section C) are based on Ghanaian labour market rates for reinforced concrete works and include formwork, steel fixing, concreting, and curing.",
        "The deterministic approach (quantity × unit price) ensures transparency and allows independent verification by project stakeholders.",
    ]
    for n in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(n)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("333333")

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")

if __name__ == "__main__":
    build_document()
